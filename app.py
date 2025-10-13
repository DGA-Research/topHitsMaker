import io
import os
import datetime
import platform
import subprocess
from typing import Tuple

import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

# -----------------------------
# Text helpers (same behavior)
# -----------------------------
def to_sentence_case(text: str) -> str:
    s = (text or "")
    lowered = ''.join(ch.lower() if ch.isalpha() else ch for ch in s)
    chars = list(lowered)
    n = len(chars)
    cap_next = True
    i = 0
    while i < n:
        ch = chars[i]
        if cap_next and ch.isalpha():
            chars[i] = ch.upper()
            cap_next = False
        elif ch in ('.', '!', '?', ':'):
            cap_next = True
        i += 1
    return ''.join(chars)

def ensure_terminal_period(s: str) -> str:
    s = (s or "").rstrip()
    if not s:
        return s
    terminal = {'.', '!', '?'}
    closing_quotes = {'"', "'", '”', '’'}
    last = s[-1]
    if last in terminal:
        return s
    if last in closing_quotes:
        if len(s) >= 2 and s[-2] in terminal:
            return s
        return s[:-1] + '.' + last
    return s + '.'

# -----------------------------
# Layout helpers
# -----------------------------
def _set_xml_indent(para, left_twips: int, hanging_twips: int):
    pPr = para._p.get_or_add_pPr()
    for ind_node in pPr.xpath('./w:ind'):
        pPr.remove(ind_node)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(left_twips))
    ind.set(qn('w:hanging'), str(hanging_twips))
    pPr.append(ind)

def _add_visible_blank_line(doc: Document):
    spacer = doc.add_paragraph("\u00A0")
    spacer.paragraph_format.line_spacing = 1.0
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)

# -----------------------------
# Styles & renderers
# -----------------------------
def ensure_h5_subbullet_style(doc: Document):
    styles = doc.styles
    if "H5Subbullet" not in [s.name for s in styles]:
        stl = styles.add_style("H5Subbullet", WD_STYLE_TYPE.PARAGRAPH)
        try:
            stl.base_style = styles["List Bullet 2"]
        except KeyError:
            stl.base_style = styles["List Bullet"]
        stl.font.name = "Arial"
        stl.font.size = Pt(10)
        pf = stl.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

def add_heading(out: Document, text: str):
    p = out.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run((text or "").upper())
    r.bold = True
    r.underline = True
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    _add_visible_blank_line(out)

def add_subheading_all_caps(out: Document, text: str):
    p = out.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run((text or "").upper())
    r.bold = False
    r.underline = True
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    _add_visible_blank_line(out)

def add_bullet(out: Document, text: str, indent: bool = False, add_period=False):
    formatted = to_sentence_case(text)
    if add_period:
        formatted = ensure_terminal_period(formatted)
    style_name = 'List Bullet 2' if indent else 'List Bullet'
    p = out.add_paragraph(formatted, style=style_name)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if p.runs:
        r = p.runs[0]
        r.font.name = 'Arial'
        r.font.size = Pt(10)
    _add_visible_blank_line(out)

def add_h5_bullet(out: Document, text: str):
    formatted = ensure_terminal_period(to_sentence_case(text))
    p = out.add_paragraph(formatted, style='H5Subbullet')
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if p.runs:
        r = p.runs[0]
        r.font.name = 'Arial'
        r.font.size = Pt(10)
    # Indent: text 0.75" (1080 twips), hanging 0.25" (360) → bullet at 0.5"
    _set_xml_indent(p, left_twips=1080, hanging_twips=360)
    _add_visible_blank_line(out)

def guess_heading_level(style_name: str, text: str) -> int:
    style = (style_name or "").lower().strip()
    t = (text or "").strip()
    if not t:
        return 0
    if "heading 2" in style or style == "h2":
        return 1
    elif "heading 3" in style or style == "h3":
        return 2
    elif "heading 4" in style or style == "h4":
        return 3
    elif "heading 5" in style or style == "h5":
        return 4
    return 0

# -----------------------------
# Core transform
# -----------------------------
def transform_docx(src: Document, add_period_to_h4: bool = True) -> Tuple[Document, dict]:
    out = Document()
    ensure_h5_subbullet_style(out)

    # Page setup
    for s in out.sections:
        s.left_margin = Inches(0.5)
        s.right_margin = Inches(0.5)

    # Base style
    base_style = out.styles['Normal']
    base_style.font.name = 'Arial'
    base_style.font.size = Pt(10)
    base_style.paragraph_format.line_spacing = 1.0
    base_style.paragraph_format.space_before = Pt(0)
    base_style.paragraph_format.space_after = Pt(0)

    counts = dict(H2=0, H3=0, H4=0, H5=0)

    for para in src.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        level = guess_heading_level(getattr(para.style, "name", ""), text)
        if level == 1:
            add_heading(out, text); counts["H2"] += 1
        elif level == 2:
            add_subheading_all_caps(out, text); counts["H3"] += 1
        elif level == 3:
            add_bullet(out, text, indent=True, add_period=add_period_to_h4); counts["H4"] += 1
        elif level == 4:
            add_h5_bullet(out, text); counts["H5"] += 1
        else:
            # Ignore non-heading text
            pass

    return out, counts

# -----------------------------
# Streamlit UI
# -----------------------------
st.set_page_config(page_title="Key Points Maker", page_icon="📝", layout="centered")
st.title("📝 Key Points Maker")
st.caption("Upload a .docx with Heading 2–5 structure. Get back a polished ‘Key Points’ document.")

uploaded = st.file_uploader("Upload your .docx", type=["docx"])
add_period_to_h4 = st.checkbox("Ensure terminal period on H4 bullets", value=True)

if uploaded is not None:
    try:
        # Load source document
        src_doc = Document(uploaded)

        # Transform
        out_doc, counts = transform_docx(src_doc, add_period_to_h4=add_period_to_h4)

        total = sum(counts.values())
        if total == 0:
            st.warning("No H2/H3/H4/H5 content found. Nothing to transform.")
        else:
            st.success(f"Transformed items — H2: {counts['H2']} • H3: {counts['H3']} • H4: {counts['H4']} • H5: {counts['H5']}")

            # Save to BytesIO
            bio = io.BytesIO()
            out_doc.save(bio)
            bio.seek(0)

            # Naming logic: if input name ends with 'INPUT' (case-insensitive), change to 'OUTPUT'.
            # Otherwise, append '_OUTPUT'.
            base_name = os.path.splitext(uploaded.name)[0].rstrip()
            if base_name.upper().endswith("INPUT"):
                out_name = base_name[:-5] + "OUTPUT.docx"  # replace trailing INPUT with OUTPUT
            else:
                out_name = base_name + "_OUTPUT.docx"

            st.download_button(
                label="⬇️ Download reformatted .docx",
                data=bio,
                file_name=out_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            st.info("Tip: In Word → Styles Pane → right-click **H5Subbullet** → ‘Select All X Instances’ → change bullet to hollow if desired.")

    except Exception as e:
        st.error(f"Error processing document: {e}")
        st.stop()

# Footer help
with st.expander("Help & notes"):
    st.markdown(
        """
- The app looks specifically for **Heading 2–5** styles in the uploaded file.
- H2 and H3 become ALL-CAPS headings; H4 becomes sentence-case bullets; H5 becomes indented sub-bullets (`H5Subbullet` style).
- Paragraphs not styled as H2–H5 are ignored on purpose.
- Everything is formatted in **Arial 10**, single-spaced, with narrow margins (0.5").
"""
    )
